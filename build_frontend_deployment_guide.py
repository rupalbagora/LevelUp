from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "LevelUp_Frontend_Vercel_Deployment_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(row.cells[idx])


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("LevelUp Frontend Deployment on Vercel")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Step-by-step setup for deploying the React/Vite frontend and updating the already-deployed backend.")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(36, 41, 47)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "F1F3F4")
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_note(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(67, 67, 67)
    r = p.add_run(body)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(67, 67, 67)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 20, "000000", 20, 6),
        ("Heading 2", 16, "000000", 18, 6),
        ("Heading 3", 14, "434343", 16, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_h1(doc, "1. Project Summary")
    add_body(
        doc,
        "Your frontend is a Vite + React app inside the existing MERN project. The backend is already deployed, so Vercel only needs to build and host the client folder. Do not deploy the whole LevelUp repository as the app root; point Vercel to the frontend folder.",
    )
    add_table(
        doc,
        ["Setting", "Value for this project"],
        [
            ["Vercel root directory", "client/levelUp"],
            ["Framework preset", "Vite"],
            ["Install command", "npm install"],
            ["Build command", "npm run build"],
            ["Output directory", "dist"],
            ["Frontend env key", "VITE_API_URL"],
            ["Google env key", "VITE_GOOGLE_CLIENT_ID"],
        ],
        [2.2, 4.0],
    )

    add_h1(doc, "2. Before You Deploy")
    add_bullets(
        doc,
        [
            "Confirm the backend deployed URL is working, for example: https://your-backend-domain.com/test.",
            "Confirm backend CORS allows the Vercel frontend domain after deployment.",
            "Make sure your frontend code is pushed to GitHub, GitLab, or Bitbucket.",
            "Keep all secrets out of Git. Add runtime values in Vercel Environment Variables.",
            "If Google sign-in is used, keep the same OAuth client ID in frontend and backend configuration.",
        ],
    )

    add_h1(doc, "3. Deploy Frontend on Vercel")
    add_numbers(
        doc,
        [
            "Open https://vercel.com and sign in.",
            "Click Add New, then Project.",
            "Import the Git repository that contains the LevelUp project.",
            "When Vercel asks for Root Directory, select client/levelUp.",
            "Set Framework Preset to Vite if Vercel does not auto-detect it.",
            "Keep Install Command as npm install.",
            "Set Build Command to npm run build.",
            "Set Output Directory to dist.",
            "Add all required environment variables before clicking Deploy.",
            "Click Deploy and wait until the build finishes successfully.",
        ],
    )

    add_h2(doc, "Environment Variables in Vercel")
    add_table(
        doc,
        ["Variable", "Example / Meaning"],
        [
            ["VITE_API_URL", "Your deployed backend base URL, e.g. https://your-backend-domain.com"],
            ["VITE_GOOGLE_CLIENT_ID", "Google OAuth Web Client ID used by @react-oauth/google"],
        ],
        [2.2, 4.0],
    )
    add_note(
        doc,
        "Important",
        "Vite exposes only variables that start with VITE_. Do not use REACT_APP_ for this project.",
    )

    add_h1(doc, "4. Backend CORS Setup for Vercel")
    add_body(
        doc,
        "After frontend deployment, copy the Vercel production URL, then add it to the backend CORS allowed origins. If the backend currently allows every origin during development, tighten it for production with your exact Vercel domain.",
    )
    add_code(
        doc,
        [
            "const allowedOrigins = [",
            "  'http://localhost:5173',",
            "  'https://your-vercel-app.vercel.app',",
            "];",
            "",
            "app.use(cors({",
            "  origin: allowedOrigins,",
            "  credentials: true,",
            "}));",
        ],
    )
    add_body(
        doc,
        "If you use cookie-based login across different domains, the backend cookie settings also matter. In production, use secure: true and sameSite: 'none' only when your frontend and backend are on different domains and HTTPS is enabled.",
    )
    add_code(
        doc,
        [
            "res.cookie('token', token, {",
            "  httpOnly: true,",
            "  secure: process.env.NODE_ENV === 'production',",
            "  sameSite: process.env.NODE_ENV === 'production' ? 'none' : 'lax',",
            "  path: '/',",
            "});",
        ],
    )

    add_h1(doc, "5. Google Sign-In Setup")
    add_numbers(
        doc,
        [
            "Open Google Cloud Console.",
            "Go to APIs & Services, then Credentials.",
            "Open your OAuth 2.0 Web Client ID.",
            "Add http://localhost:5173 to Authorized JavaScript origins for local testing.",
            "Add your Vercel URL, for example https://your-vercel-app.vercel.app, to Authorized JavaScript origins.",
            "Save the changes.",
            "Use the same client ID in Vercel as VITE_GOOGLE_CLIENT_ID.",
            "Use the same client ID in the backend deployment as GOOGLE_CLIENT_ID.",
        ],
    )

    add_h1(doc, "6. Test After Deployment")
    add_bullets(
        doc,
        [
            "Open the Vercel URL in a fresh browser session.",
            "Test normal signup with username, email, and password.",
            "Test login and confirm the dashboard opens.",
            "Test forgot password: request token, reset password, then login with the new password.",
            "Test Google sign-in after OAuth origins are updated.",
            "Open Profile and Dashboard pages to confirm authenticated API calls work.",
            "Create a battle link and verify the frontend can communicate with the deployed backend.",
            "Check Vercel Function/Build logs only for frontend build issues; backend runtime logs will be on the backend hosting platform.",
        ],
    )

    add_h1(doc, "7. Common Errors and Fixes")
    add_table(
        doc,
        ["Problem", "Most likely fix"],
        [
            ["Blank page on Vercel", "Check Root Directory is client/levelUp and Output Directory is dist."],
            ["API calls go to wrong URL", "Set VITE_API_URL in Vercel and redeploy."],
            ["Login works locally but not on Vercel", "Update backend CORS and cookie sameSite/secure settings."],
            ["Google button fails", "Add Vercel URL to Google OAuth Authorized JavaScript origins."],
            ["Environment variable not found", "Use VITE_ prefix and redeploy after adding the variable."],
            ["404 on refresh", "Vercel usually handles Vite SPA routing; add a rewrite to /index.html if needed."],
        ],
        [2.2, 4.0],
    )

    add_h2(doc, "Optional vercel.json for SPA refresh")
    add_body(doc, "Create this file inside client/levelUp only if refresh routes like /dashboard or /signin show 404.")
    add_code(
        doc,
        [
            "{",
            '  "rewrites": [',
            '    { "source": "/(.*)", "destination": "/index.html" }',
            "  ]",
            "}",
        ],
    )

    add_h1(doc, "8. How to Add New Backend Features to the Already-Deployed Backend")
    add_body(
        doc,
        "When you add a new backend feature, the deployed backend will not update automatically unless your hosting platform is connected to Git auto-deploy or you manually redeploy it. Follow this workflow.",
    )
    add_numbers(
        doc,
        [
            "Add the backend code locally: model, controller, route, middleware, or service as needed.",
            "Register the route in server/src/app.js or the correct route file.",
            "Add any new environment variables to the backend hosting platform dashboard.",
            "Run backend tests locally: npm.cmd test from LevelUp/server.",
            "Commit and push the backend changes to the branch used by your backend host.",
            "Open your backend hosting platform and trigger redeploy if auto-deploy is not enabled.",
            "Check deployment logs for install, build, startup, MongoDB, and env errors.",
            "Test the new API endpoint using Postman, Thunder Client, or the frontend.",
            "If the frontend calls this new API, update frontend service files and redeploy Vercel too.",
        ],
    )
    add_note(
        doc,
        "Rule",
        "Backend feature added means backend redeploy is required. Frontend redeploy is required only when frontend code or frontend environment variables changed.",
    )

    add_h1(doc, "9. Final Deployment Checklist")
    add_bullets(
        doc,
        [
            "Frontend root directory on Vercel is client/levelUp.",
            "VITE_API_URL points to deployed backend, not localhost.",
            "VITE_GOOGLE_CLIENT_ID is set in Vercel.",
            "Backend GOOGLE_CLIENT_ID is set on backend host.",
            "Backend CORS includes the Vercel production URL.",
            "Cookie settings are production-ready if using cross-domain cookies.",
            "Signup, login, forgot password, reset password, Google sign-in, dashboard, and profile are tested after deploy.",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
